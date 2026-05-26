#!/usr/bin/env python3
"""
Generate a plain-language client briefing paper as .docx.

This script is a worked example from the BMM/Worner engagement analysis.
Adapt it per engagement: swap the title, obligations table, timeline,
and issue descriptions.

Usage:
    python3 generate_briefing_docx.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- TITLE ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CONFIDENTIAL')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(180, 0, 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n[PROJECT NAME]\nBriefing Paper for [CLIENT NAMES]')
run.bold = True
run.font.size = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nPrepared: [DATE]')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# --- SECTION STRUCTURE (fill in per engagement) ---
# Part 1: What [Provider] Owed the Client
# Part 2: What [Provider] Actually Did
# Part 3: Timeline
# Part 4: Rights, Breaches, and Options
# Part 5: Legal Concepts in Simple Terms
# Summary of Recommendations

# See the parent SKILL.md for detailed structure guidance
# Key: write as "briefing a client in simple terms"
